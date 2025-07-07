import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown_content: str) -> bytes:
    """
    将 Markdown 内容转换为 Word 文档
    
    Args:
        markdown_content: Markdown 格式的文本内容
        
    Returns:
        Word 文档的字节数据
    """
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # 空行
            doc.add_paragraph()
            continue
            
        # 处理标题
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('##### '):
            heading = doc.add_heading(line[6:], level=5)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('###### '):
            heading = doc.add_heading(line[7:], level=6)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # 普通段落
            paragraph = doc.add_paragraph()
            
            # 处理粗体、斜体和代码
            text_parts = []
            current_text = line
            
            # 处理粗体
            bold_pattern = r'\*\*(.*?)\*\*'
            bold_matches = list(re.finditer(bold_pattern, current_text))
            
            # 处理斜体
            italic_pattern = r'\*(.*?)\*'
            italic_matches = list(re.finditer(italic_pattern, current_text))
            
            # 处理代码
            code_pattern = r'`(.*?)`'
            code_matches = list(re.finditer(code_pattern, current_text))
            
            # 合并所有匹配项并按位置排序
            all_matches = []
            for match in bold_matches:
                all_matches.append((match.start(), match.end(), 'bold', match.group(1)))
            for match in italic_matches:
                all_matches.append((match.start(), match.end(), 'italic', match.group(1)))
            for match in code_matches:
                all_matches.append((match.start(), match.end(), 'code', match.group(1)))
            
            all_matches.sort(key=lambda x: x[0])
            
            # 构建段落
            last_end = 0
            for start, end, style, content in all_matches:
                # 添加前面的普通文本
                if start > last_end:
                    normal_text = current_text[last_end:start]
                    if normal_text:
                        paragraph.add_run(normal_text)
                
                # 添加格式化的文本
                run = paragraph.add_run(content)
                if style == 'bold':
                    run.bold = True
                elif style == 'italic':
                    run.italic = True
                elif style == 'code':
                    run.font.name = 'Courier New'
                
                last_end = end
            
            # 添加剩余的普通文本
            if last_end < len(current_text):
                remaining_text = current_text[last_end:]
                if remaining_text:
                    paragraph.add_run(remaining_text)
    
    # 将文档保存到字节流
    from io import BytesIO
    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    return docx_stream.getvalue()


def generate_report_filename() -> str:
    """
    生成报告文件名
    
    Returns:
        格式化的文件名
    """
    now = datetime.now()
    timestamp = now.strftime("%Y-%m-%d_%H-%M-%S")
    return f"research-report-{timestamp}.docx" 